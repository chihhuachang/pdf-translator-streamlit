import streamlit as st
import google.generativeai as genai
import os
import sys
import io
import time # 引入 time 模組用於添加延遲

# Import libraries for file reading AND writing docx
import docx
from docx import Document
import PyPDF2
from PyPDF2.errors import PdfReadError

# --- 常數定義 ---
# 定義每個文本塊的最大字符數 (可根據需要調整， Gemeni Pro 約有 30k token 限制，Flash 較少，保守一點)
# 注意：這只是個大概值，實際 token 數與字符數不完全對應
MAX_CHARS_PER_CHUNK = 2500
# 每次 API 呼叫之間的延遲（秒），避免觸及速率限制 (免費版通常有每分鐘請求限制)
API_CALL_DELAY = 2

# --- 說明文件 ---
# 功能：上傳英文文件(txt, docx, pdf)，讀取內容，將長文本 **分塊(Chunking)**，
#       使用固定的詳細提示詞和流式傳輸逐塊翻譯成繁體中文，
#       合併結果後提供 .docx 檔案下載。
# (其餘說明與之前版本相同)
#
# 如何執行：
# 1. 安裝函式庫: pip install streamlit google-generativeai python-docx PyPDF2
# 2. 將此程式碼儲存為 `translate_file_app_en_to_zh_chunking.py`。
# 3. 在終端機中執行： streamlit run translate_file_app_en_to_zh_chunking.py
# ---

# --- API 金鑰設定 (保持不變) ---
api_key = os.getenv("GOOGLE_API_KEY")
# ...(省略與之前版本相同的 API Key 檢查與設定程式碼)...
if not api_key: st.error("..."); st.stop()
try: genai.configure(api_key=api_key)
except Exception as e: st.error(f"...: {e}"); st.stop()


# --- 從檔案提取文字的函式 (保持不變) ---
def extract_text_from_file(uploaded_file):
    """
    根據上傳檔案的類型提取文字內容。
    支援 .txt, .docx, .pdf 格式。
    (此函式邏輯與上一版本完全相同)
    """
    extracted_text = ""
    # ...(省略與上一版本相同的檔案讀取和錯誤處理邏輯)...
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension == ".txt":
            try: extracted_text = uploaded_file.getvalue().decode("utf-8")
            except UnicodeDecodeError:
                st.warning("嘗試 UTF-8 解碼失敗，嘗試 Big5...")
                try: extracted_text = uploaded_file.getvalue().decode("big5", errors='ignore')
                except Exception as e_enc:
                     st.error(f"嘗試 Big5 解碼也失敗: {e_enc}。使用忽略錯誤的 UTF-8。")
                     extracted_text = uploaded_file.getvalue().decode("utf-8", errors='ignore')
            st.info(f"成功讀取 .txt 檔案: {uploaded_file.name}")
        elif file_extension == ".docx":
            document = docx.Document(uploaded_file)
            extracted_text = '\n'.join([para.text for para in document.paragraphs])
            st.info(f"成功讀取 .docx 檔案: {uploaded_file.name}")
        elif file_extension == ".pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                if pdf_reader.is_encrypted: st.error("錯誤：PDF 文件已加密。"); return None
                full_text = [page.extract_text() for i, page in enumerate(pdf_reader.pages) if page.extract_text() or st.warning(f"讀取 PDF 第 {i+1} 頁時未提取到文字或發生錯誤。", icon="⚠️")]
                extracted_text = '\n'.join(filter(None, full_text))
                if not extracted_text.strip(): st.warning(f"無法從 PDF '{uploaded_file.name}' 提取任何文字。")
                else: st.info(f"成功讀取 .pdf 檔案: {uploaded_file.name}")
            except PdfReadError as pdf_err: st.error(f"PyPDF2 讀取錯誤: {pdf_err}"); return None
        else: st.error(f"錯誤：不支援的檔案類型 '{file_extension}'。"); return None
        return extracted_text.strip()
    except Exception as e: st.error(f"讀取或解析檔案 '{uploaded_file.name}' 時發生錯誤: {e}"); return None


# --- 新增：文本分塊函式 ---
def split_text_into_chunks(text, max_chars=MAX_CHARS_PER_CHUNK):
    """
    將長文本分割成多個塊，盡量按段落分割，並確保每塊不超過最大字符數。
    """
    chunks = []
    # 先嘗試按段落分割 (假設段落由兩個換行符分隔)
    paragraphs = text.split('\n\n')
    current_chunk = ""

    for paragraph in paragraphs:
        # 如果段落本身就超長，需要強制分割 (這裡用簡單的字符分割，可改進為按句子)
        if len(paragraph) > max_chars:
            # 強制分割長段落
            start = 0
            while start < len(paragraph):
                # 嘗試找到一個合理的斷點 (例如句號、問號後) 靠近 max_chars
                # 簡化處理：直接按 max_chars 切割
                end = min(start + max_chars, len(paragraph))
                long_para_chunk = paragraph[start:end]

                # 檢查加上這部分是否會讓 current_chunk 超長
                if len(current_chunk) + len(long_para_chunk) + 2 <= max_chars: # +2 for '\n\n'
                     if current_chunk:
                         current_chunk += "\n\n" + long_para_chunk
                     else:
                         current_chunk = long_para_chunk
                else:
                    # 如果 current_chunk 已經有內容，先儲存它
                    if current_chunk:
                        chunks.append(current_chunk)
                    # 這個長段落的片段自成一塊 (或成為下一塊的開頭)
                    current_chunk = long_para_chunk
                    # 如果這個片段本身就等於最大長度，立刻存檔開始新的
                    if len(current_chunk) >= max_chars:
                        chunks.append(current_chunk)
                        current_chunk = "" # 開始新的 chunk

                start = end # 移動到下一個切割點
            # 長段落處理完後，可能還有剩餘的 current_chunk 需要處理
            if current_chunk: # 如果長段落切完最後一部分成為了新的 current_chunk
                 pass # 繼續處理下一個段落，看能否合併

        # 如果段落不超長
        else:
            # 檢查加上這個段落是否會讓 current_chunk 超長
            if len(current_chunk) + len(paragraph) + 2 <= max_chars: # +2 for '\n\n'
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                # 當前塊已滿，儲存當前塊，並用此段落開始新塊
                chunks.append(current_chunk)
                current_chunk = paragraph

    # 添加最後一個塊 (如果非空)
    if current_chunk:
        chunks.append(current_chunk)

    return chunks


# --- 翻譯函式 (保持使用流式傳輸) ---
def translate_text(text_to_translate, target_language="繁體中文"):
    """
    使用 Gemini 模型和固定的詳細提示詞，透過流式傳輸翻譯單個文本塊。
    (此函式邏輯與上一版本完全相同)
    """
    if not text_to_translate: return None # 空塊直接返回

    # --- 使用簡化的固定提示詞 ---
    # 替換掉原本詳細的多步驟提示詞
    fixed_instruction_prompt = "Please translate the following English text into accurate and natural Traditional Chinese:"
    # fixed_instruction_prompt = """...""".strip() # 省略，使用你之前提供的完整提示詞
    # --- 使用你提供的固定提示詞 ---
    # fixed_instruction_prompt = """
    # Please act as a professional Chinese translator.
    # I will give you a piece of text, and you will actually follow the steps below to produce a professional Chinese translation that satisfies me.
    # 1. Carefully read and fully understand the original text, ensuring thorough comprehension without haste.
    # 2. Carefully think and consider how you would share the content you just read with your imagined audience in Chinese.
    # 3. Start to translate the text by writting down the proposed sharing content you just had with your imagined audience using traditional Chinese characters. Avoid translating word-for-word; aim for a comfortable, natural, and smooth manner of expression.
    # 4. Compare it with the original text to identify any omissions or inaccuracies, then refine as necessary.
    # """.strip()

    # --- 組合提示詞和待翻譯文字 (保持不變) ---
    full_prompt = f"{fixed_instruction_prompt}\n\n{text_to_translate}"
    
    # 選擇模型 (保持不變)
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
    # model = genai.GenerativeModel('gemini-2.0-flash')

    # 狀態提示 (保持不變)
    # st.info("正在使用流式傳輸將英文文字傳送給模型進行翻譯...") # 在迴圈中顯示太頻繁

    try:
        response = model.generate_content(full_prompt, stream=True)
        full_translated_text = "".join(chunk.text for chunk in response if hasattr(chunk, 'text') and chunk.text)
        return full_translated_text.strip()
    except Exception as e:
        st.error(f"翻譯塊時 API 呼叫或流式處理過程中發生錯誤: {e}")
        return f"[[翻譯錯誤: {e}]]" # 返回錯誤標記而非 None，以便追蹤


# --- 從文字建立 Docx 檔案函式 (保持不變) ---
def create_docx_from_text(text_content, base_filename):
    """
    將文字內容儲存為 .docx 檔案的 BytesIO 物件。
    (此函式邏輯與上一版本完全相同)
    """
    try:
        document = Document()
        # 按段落添加，保留換行
        for paragraph in text_content.split('\n'):
            document.add_paragraph(paragraph)
        # document.add_paragraph(text_content) # 原本的方式
        docx_buffer = io.BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        docx_filename = f"{base_filename}_translated_chunked.docx" # 修改檔名區別
        return {'name': docx_filename, 'data': docx_buffer}, None
    except Exception as e:
        error_message = f"建立 Docx 檔案時發生錯誤: {e}"
        st.error(error_message)
        return {'name': None, 'data': None}, error_message


# --- Streamlit 應用程式介面 (略作調整) ---
st.set_page_config(page_title="文件分塊翻譯+下載 (英->繁中)", layout="wide")
st.title("📝 Gemini 文件分塊翻譯 (英文 ➔ 繁體中文) 並下載 Docx") # 修改
st.caption(f"上傳 .txt, .docx, 或 .pdf 英文文件，將文本分塊 (每塊約 {MAX_CHARS_PER_CHUNK} 字符) 翻譯後合併並提供 Docx 下載") # 修改

# ...(省略與之前版本相同的 col1, col2, uploaded_file, base_filename, translate_button 定義)...
col1, col2 = st.columns(2)
uploaded_file = None
base_filename = None
with col1:
    st.subheader("步驟 1: 上傳英文文件")
    uploaded_file = st.file_uploader(
        "選擇要翻譯的英文文件 (.txt, .docx, .pdf)", type=['txt', 'docx', 'pdf'], key="file_uploader"
    )
    if uploaded_file is not None:
        st.markdown(f"**已上傳檔案:** `{uploaded_file.name}` (`{uploaded_file.type}`)")
        base_filename = os.path.splitext(uploaded_file.name)[0]
    translate_button = st.button("開始分塊翻譯成繁體中文", key="translate_btn", disabled=uploaded_file is None) # 修改按鈕文字

with col2:
    st.subheader("翻譯結果 (繁體中文)")
    result_placeholder = st.empty()
    result_placeholder.text_area(
        label="翻譯結果預覽區域", value="翻譯結果將會顯示在這裡...", height=400,
        key="result_text_area", disabled=True, label_visibility="collapsed"
    )
    download_placeholder = st.empty()
    # 新增進度條佔位符
    progress_placeholder = st.empty()


# --- 執行文件讀取、分塊、翻譯、合併 (主要修改部分) ---
if translate_button:
    if uploaded_file is None:
        st.warning("請先上傳一個文件。")
    else:
        download_placeholder.empty() # 清空舊下載按鈕
        progress_placeholder.empty() # 清空舊進度條
        result_placeholder.text_area(
            label="翻譯結果預覽區域", value="處理中...", height=400,
            key="result_text_area_processing", disabled=True, label_visibility="collapsed"
        )

        # 步驟 1: 讀取文件
        with st.spinner(f"正在讀取檔案 '{uploaded_file.name}'..."):
            extracted_text = extract_text_from_file(uploaded_file)

        if extracted_text is not None and extracted_text.strip():
            st.success("成功從文件中提取英文文字！")

            # 步驟 2: 將提取的文字分塊
            with st.spinner("正在將文本分割成處理塊..."):
                text_chunks = split_text_into_chunks(extracted_text, MAX_CHARS_PER_CHUNK)
                total_chunks = len(text_chunks)
                if total_chunks == 0:
                    st.warning("未能將文本有效分割成塊。")
                    st.stop()
                st.info(f"文本已分割成 {total_chunks} 個塊進行翻譯。")

            # 步驟 3: 逐塊翻譯並合併結果
            translated_chunks = []
            errors_occurred = False
            progress_bar = progress_placeholder.progress(0) # 初始化進度條
            status_text = progress_placeholder.text(f"正在翻譯塊 1 / {total_chunks}...")

            fixed_target_language = "繁體中文"

            for i, chunk in enumerate(text_chunks):
                chunk_num = i + 1
                status_text.text(f"正在翻譯塊 {chunk_num} / {total_chunks}...")
                # (可選) 顯示正在處理的塊 (用於除錯)
                # with st.expander(f"查看塊 {chunk_num} 原文 (部分)"):
                #    st.text(chunk[:200] + "...")

                # 呼叫翻譯函式 (內部使用流式)
                translated_chunk = translate_text(chunk, fixed_target_language)

                if translated_chunk and "[[翻譯錯誤:" not in translated_chunk:
                    translated_chunks.append(translated_chunk)
                else:
                    errors_occurred = True
                    translated_chunks.append(f"\n--- 塊 {chunk_num} 翻譯失敗 ---\n{translated_chunk or '未知錯誤'}\n---") # 保留錯誤訊息
                    st.error(f"翻譯塊 {chunk_num} 時發生錯誤。")

                # 更新進度條
                progress_bar.progress(chunk_num / total_chunks)

                # 在兩次 API 呼叫之間加入延遲
                if chunk_num < total_chunks:
                    time.sleep(API_CALL_DELAY)

            status_text.text("合併翻譯結果...")

            # 步驟 4: 合併翻譯結果
            # 使用換行符連接，如果原始分割是按段落，這會比較自然
            final_translated_text = "\n\n".join(translated_chunks)

            # 步驟 5: 顯示結果和提供下載
            result_placeholder.text_area(
                label="翻譯結果 (繁體中文) - 合併自塊:", value=final_translated_text, height=400,
                key="result_text_area_updated", disabled=False, label_visibility="visible"
            )

            if errors_occurred:
                st.warning("部分文本塊翻譯失敗，請檢查上方錯誤訊息及合併結果中的標記。")
            else:
                st.success("所有文本塊翻譯完成！")

            # 產生並提供 Docx 下載 (使用合併後的文本)
            with st.spinner("正在產生 Docx 檔案..."):
                docx_data, docx_error = create_docx_from_text(final_translated_text, base_filename)
            if not docx_error and docx_data['data']:
                download_placeholder.download_button(
                    label=f"📥 下載合併翻譯結果 ({docx_data['name']})",
                    data=docx_data['data'], file_name=docx_data['name'],
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key='download_docx'
                )
            else:
                download_placeholder.error("抱歉，產生 Docx 檔案時發生錯誤。")

        # ... (處理提取文字失敗或為空的情況，保持不變) ...
        elif extracted_text is not None and not extracted_text.strip():
             st.warning("從檔案中未提取到任何有效文字內容，無法進行翻譯。")
             result_placeholder.text_area("翻譯結果 (繁體中文):", value="未提取到文字。", height=400, key="result_text_area_no_text", disabled=False, label_visibility="visible")
        else:
             st.error("無法從文件中提取文字，請檢查檔案格式或查看上方的錯誤訊息。")
             result_placeholder.text_area("翻譯結果 (繁體中文):", value="文件讀取或文字提取失敗。", height=400, key="result_text_area_extract_fail", disabled=False, label_visibility="visible")